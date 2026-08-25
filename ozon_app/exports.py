"""Local report and replenishment spreadsheet exports."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable


class ExportError(RuntimeError):
    """Raised when a local report cannot be generated."""


def _target(path: Path, suffix: str) -> Path:
    value = Path(path)
    if value.suffix.lower() != suffix:
        value = value.with_suffix(suffix)
    value.parent.mkdir(parents=True, exist_ok=True)
    return value


def _markdown_lines(text: str) -> Iterable[tuple[str, str]]:
    for source in str(text or "").replace("\r\n", "\n").split("\n"):
        line = source.strip()
        if not line:
            yield "blank", ""
        elif line.startswith("## "):
            yield "heading2", line[3:].strip()
        elif line.startswith("# "):
            yield "heading1", line[2:].strip()
        elif re.match(r"^[-*]\s+", line):
            yield "bullet", re.sub(r"^[-*]\s+", "", line)
        elif re.match(r"^\d+[.)]\s+", line):
            yield "number", re.sub(r"^\d+[.)]\s+", "", line)
        else:
            yield "paragraph", re.sub(r"\*\*(.+?)\*\*", r"\1", line)


def export_ai_docx(
    path: Path, report_text: str, date_from: str, date_to: str, model: str,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as error:
        raise ExportError("Word 导出组件不可用，请安装 python-docx。") from error
    target = _target(path, ".docx")
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    document.add_heading("Ozon AI 经营分析报告", level=0)
    document.add_paragraph(f"分析周期：{date_from} 至 {date_to}")
    document.add_paragraph(f"分析模型：{model or '未命名模型'}")
    document.add_paragraph(
        f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    for kind, line in _markdown_lines(report_text):
        if kind == "blank":
            continue
        if kind == "heading1":
            document.add_heading(line, level=1)
        elif kind == "heading2":
            document.add_heading(line, level=2)
        elif kind == "bullet":
            document.add_paragraph(line, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(line, style="List Number")
        else:
            document.add_paragraph(line)
    try:
        document.save(target)
    except OSError as error:
        raise ExportError(f"无法写入 Word 报告：{error}") from error
    return target


def export_ai_pdf(
    path: Path, report_text: str, date_from: str, date_to: str, model: str,
) -> Path:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as error:
        raise ExportError("PDF 导出组件不可用，请安装 reportlab。") from error
    target = _target(path, ".pdf")
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "OzonTitle", parent=styles["Title"], fontName="Helvetica-Bold",
        textColor=colors.HexColor("#175CD3"), spaceAfter=12,
    )
    heading = ParagraphStyle(
        "OzonHeading", parent=styles["Heading2"], fontName="Helvetica-Bold",
        textColor=colors.HexColor("#172B4D"), spaceBefore=9, spaceAfter=5,
    )
    body = ParagraphStyle(
        "OzonBody", parent=styles["BodyText"], fontName="Helvetica",
        leading=16, spaceAfter=5,
    )
    story: list[Any] = [
        Paragraph("Ozon AI 经营分析报告", title),
        Paragraph(f"分析周期：{date_from} 至 {date_to}<br/>分析模型：{model or '未命名模型'}", body),
        Spacer(1, 0.25 * cm),
    ]
    for kind, line in _markdown_lines(report_text):
        if kind == "blank":
            story.append(Spacer(1, 0.12 * cm))
        elif kind in {"heading1", "heading2"}:
            story.append(Paragraph(line, heading))
        elif kind == "bullet":
            story.append(Paragraph("• " + line, body))
        elif kind == "number":
            story.append(Paragraph(line, body))
        else:
            story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body))
    try:
        SimpleDocTemplate(
            str(target), pagesize=A4, rightMargin=1.7 * cm, leftMargin=1.7 * cm,
            topMargin=1.7 * cm, bottomMargin=1.7 * cm,
        ).build(story)
    except OSError as error:
        raise ExportError(f"无法写入 PDF 报告：{error}") from error
    return target


def export_replenishment_xlsx(
    path: Path,
    rows: list[dict[str, Any]],
    product: dict[str, Any],
    target_days: int,
    supply_type: str,
    dropoff: dict[str, Any] | None,
) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError as error:
        raise ExportError("Excel 导出组件不可用，请安装 openpyxl。") from error
    target = _target(path, ".xlsx")
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "补货计划"
    sheet.append(["Ozon 补货计划"])
    sheet.append([
        "商品",
        " / ".join(str(product.get(key) or "") for key in ("offer_id", "sku", "product_name")).strip(" / "),
    ])
    sheet.append(["供货方式", supply_type or ""])
    sheet.append([
        "发货地址",
        " / ".join(str((dropoff or {}).get(key) or "") for key in ("city", "address")).strip(" / "),
    ])
    headers = [
        "Ozon SKU", "货号", "商品名称", "宏区集群 ID", "集群中文名",
        "可售库存", "在途", "已申请", "日均销量", "目标库存（按目标天数）",
        "目标天数", "建议补货", "计划补货", "计划后可售天数",
    ]
    sheet.append(headers)
    header_row = sheet.max_row
    for index, row in enumerate(rows, start=header_row + 1):
        available = float(row.get("available_stock") or 0)
        transit = float(row.get("transit_stock") or 0)
        requested = float(row.get("requested_stock") or 0)
        daily_sales = float(row.get("daily_sales") or 0)
        planned = int(row.get("planned_qty") or 0)
        sheet.append([
            product.get("sku") or "", product.get("offer_id") or "", product.get("product_name") or "",
            row.get("macrolocal_cluster_id") or "", row.get("cluster_name") or "",
            available, transit, requested, daily_sales,
            f"=I{index}*{int(target_days or 0)}",
            int(target_days or 0),
            f"=MAX(0,ROUNDUP(I{index}*K{index}-F{index}-G{index}-H{index},0))",
            planned,
            f'=IF(I{index}>0,(F{index}+G{index}+H{index}+M{index})/I{index},"")',
        ])
    title_fill = PatternFill("solid", fgColor="EAF2FF")
    header_fill = PatternFill("solid", fgColor="2F6FED")
    sheet["A1"].font = Font(bold=True, size=15, color="175CD3")
    for cell in sheet[header_row]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center")
    for row in sheet.iter_rows(min_row=1, max_row=header_row - 1, min_col=1, max_col=2):
        for cell in row:
            cell.fill = title_fill
    for column, width in enumerate((14, 16, 24, 16, 22, 12, 12, 12, 12, 18, 12, 14, 14, 18), start=1):
        sheet.column_dimensions[get_column_letter(column)].width = width
    sheet.freeze_panes = f"A{header_row + 1}"
    sheet.auto_filter.ref = f"A{header_row}:N{max(header_row, sheet.max_row)}"
    workbook.calculation.fullCalcOnLoad = True
    try:
        workbook.save(target)
        # openpyxl escapes CJK characters as numeric entities in inline strings.
        # Keep UTF-8 labels literal for compatibility with the existing template
        # validation and with lightweight spreadsheet parsers used by operators.
        import re
        from zipfile import ZIP_DEFLATED, ZipFile

        temporary = target.with_name(target.name + ".utf8")
        with ZipFile(target, "r") as source, ZipFile(temporary, "w", ZIP_DEFLATED) as output:
            for item in source.infolist():
                payload = source.read(item.filename)
                if item.filename == "xl/worksheets/sheet1.xml":
                    payload = re.sub(
                        rb"&#(\d+);",
                        lambda match: chr(int(match.group(1))).encode("utf-8"),
                        payload,
                    )
                output.writestr(item, payload)
        temporary.replace(target)
    except OSError as error:
        raise ExportError(f"无法写入 Excel 补货计划：{error}") from error
    return target
